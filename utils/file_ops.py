"""
file_ops.py
Utilities for file validation, conversion, merging, and extraction for grading workflows.
"""

import os
import base64
import time
from pathlib import Path
from PyPDF2 import PdfReader, PdfMerger
import pytesseract
from pdf2image import convert_from_path
from PIL import Image

# Configure pytesseract path for containerized environments
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

# Optional DOCX → PDF conversion
try:
    from docx2pdf import convert
except ImportError:
    convert = None

import subprocess
import tempfile

# Optional DOCX text extraction
try:
    from docx import Document
except ImportError:
    Document = None

def check_libreoffice_available():
    """Check if LibreOffice is available on the system."""
    try:
        result = subprocess.run(['libreoffice', '--version'], 
                              capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            print(f"✅ LibreOffice available: {result.stdout.strip()}")
            return True
        else:
            print(f"❌ LibreOffice not available: {result.stderr}")
            return False
    except Exception as e:
        print(f"❌ Error checking LibreOffice: {e}")
        return False

def extract_text_from_docx(docx_path):
    """Extract text directly from a DOCX file using python-docx."""
    if Document is None:
        print("❌ python-docx not available for direct DOCX text extraction")
        return ""
    
    try:
        print(f"🔍 Extracting text directly from DOCX: {docx_path}")
        doc = Document(docx_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text.strip() + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text.strip() + " "
                text += "\n"
        
        if text.strip():
            print(f"✅ Extracted {len(text.strip())} characters from DOCX")
            return text.strip()
        else:
            print("⚠️ No text found in DOCX file")
            return ""
            
    except Exception as e:
        print(f"❌ Failed to extract text from DOCX: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    """Extracts all text from a PDF file, using OCR if needed."""
    print(f"🔍 Extracting text from {pdf_path}...")
    
    # First try PyPDF2
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text += page_text.strip() + "\n"
                print(f"📄 Page {page_num + 1}: Extracted {len(page_text.strip())} characters")
        
        if text.strip():
            print(f"✅ PyPDF2 extracted {len(text.strip())} characters total")
            return text.strip()
        else:
            print("⚠️ PyPDF2 found no text, trying OCR...")
    except Exception as e:
        print(f"❌ PyPDF2 failed for {pdf_path}: {e}")

    # Fallback to OCR if no text was extracted
    try:
        print(f"🔍 Running OCR on {pdf_path}...")
        
        # Test if tesseract is working
        try:
            pytesseract.get_tesseract_version()
            print(f"✅ Tesseract version: {pytesseract.get_tesseract_version()}")
        except Exception as e:
            print(f"❌ Tesseract not found: {e}")
            return ""
        
        images = convert_from_path(pdf_path, dpi=300)  # Higher DPI for better OCR
        print(f"📄 Converted PDF to {len(images)} image(s)")
        
        ocr_text = ""
        
        for i, image in enumerate(images):
            print(f"🔍 Processing page {i + 1} with OCR...")
            # Use better OCR settings
            page_text = pytesseract.image_to_string(
                image, 
                config='--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,!?;:()[]{}"\' '
            )
            if page_text and page_text.strip():
                ocr_text += page_text.strip() + "\n"
                print(f"📄 OCR Page {i + 1}: Extracted {len(page_text.strip())} characters")
        
        if ocr_text.strip():
            print(f"✅ OCR extracted {len(ocr_text.strip())} characters total")
            return ocr_text.strip()
        else:
            print("⚠️ OCR found no text")
            return ""
            
    except Exception as e:
        print(f"❌ OCR failed for {pdf_path}: {e}")
        import traceback
        print(f"❌ Full error traceback: {traceback.format_exc()}")
        return ""

def encode_file_to_base64(filepath):
    """Encodes a PDF or DOCX file as base64 string."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in [".pdf", ".docx"]:
        raise ValueError(f"❌ Unsupported file type: {ext}")
    with open(filepath, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def is_valid_pdf(filepath):
    """Returns True if PDF is valid (not encrypted, not empty, readable), else False."""
    try:
        reader = PdfReader(filepath)
        if reader.is_encrypted:
            print(f"⚠️ Skipping encrypted PDF: {filepath}")
            return False
        if len(reader.pages) == 0:
            print(f"⚠️ PDF has no pages: {filepath}")
            return False
        # Don't require extractable text - some PDFs are image-based but still valid
        return True
    except Exception as e:
        print(f"⚠️ Error reading PDF {filepath}: {e}")
        return False

def is_supported_file(filepath):
    """Returns True if the file is a supported (valid) PDF or a non-empty DOCX."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return is_valid_pdf(filepath)
    elif ext == ".docx":
        return os.path.getsize(filepath) > 0
    return False

def convert_docx_to_pdf(docx_path, max_retries=3, wait_seconds=1):
    """
    Converts a DOCX to PDF using LibreOffice, retrying up to max_retries if needed.
    Returns the path to the PDF, or raises if it fails.
    """
    if not docx_path.lower().endswith(".docx"):
        raise ValueError("Only .docx files can be converted.")
    
    pdf_path = docx_path.replace(".docx", ".pdf")
    
    # Check if LibreOffice is available
    if not check_libreoffice_available():
        print("⚠️ LibreOffice not available, trying docx2pdf fallback...")
        # Skip to fallback method
        attempt = max_retries
    else:
        attempt = 0

    while attempt < max_retries:
        try:
            print(f"🔧 Attempt {attempt + 1}: Converting DOCX to PDF using LibreOffice: {docx_path}")
            
            # Use LibreOffice headless mode to convert DOCX to PDF
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path),
                docx_path
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,  # 60 second timeout
                cwd=os.path.dirname(docx_path)
            )
            
            if result.returncode == 0:
                # Wait a moment for file to be written
                for _ in range(10):
                    if os.path.exists(pdf_path):
                        break
                    time.sleep(0.5)
                
                if os.path.exists(pdf_path) and is_valid_pdf(pdf_path):
                    print(f"✅ PDF successfully created and validated at: {pdf_path}")
                    return pdf_path
                else:
                    print(f"⚠️ PDF file created but invalid or empty: {pdf_path}")
            else:
                print(f"❌ LibreOffice conversion failed with return code {result.returncode}")
                print(f"❌ Error output: {result.stderr}")
                
        except subprocess.TimeoutExpired:
            print(f"❌ Attempt {attempt + 1} timed out for {docx_path}")
        except Exception as e:
            print(f"❌ Attempt {attempt + 1} failed for {docx_path}: {e}")
        
        attempt += 1
        if attempt < max_retries:
            time.sleep(wait_seconds)

    # Fallback to docx2pdf if LibreOffice fails and it's available
    if convert is not None:
        print(f"🔄 Falling back to docx2pdf for {docx_path}")
        try:
            import pythoncom
            pythoncom.CoInitialize()
            convert(docx_path, pdf_path)
            for _ in range(5):
                if os.path.exists(pdf_path):
                    break
                time.sleep(0.5)
            if os.path.exists(pdf_path) and is_valid_pdf(pdf_path):
                print(f"✅ PDF successfully created with docx2pdf fallback at: {pdf_path}")
                return pdf_path
        except Exception as e:
            print(f"❌ docx2pdf fallback also failed: {e}")
        finally:
            try:
                pythoncom.CoUninitialize()
            except:
                pass

    raise RuntimeError(f"❌ Failed to convert DOCX to PDF after {max_retries} attempts: {docx_path}")

def merge_pdfs(pdf_paths, output_path):
    """Merges multiple PDF files into a single PDF at output_path."""
    merger = PdfMerger()
    for path in pdf_paths:
        if path.lower().endswith(".pdf") and os.path.exists(path):
            merger.append(path)
        else:
            print(f"⚠️ Skipping non-PDF or missing file: {path}")
    merger.write(output_path)
    merger.close()

def prepare_submission_for_grading(file_paths, output_path):
    """
    Given a list of file paths (PDFs or DOCX), merges/converts all to a single PDF,
    extracts the combined text, and returns (merged_pdf_path, extracted_text).
    """
    pdfs_to_merge = []
    skipped = []
    docx_texts = []  # Store text extracted directly from DOCX files

    for path in file_paths:
        ext = Path(path).suffix.lower()
        if ext == ".docx":
            try:
                converted = convert_docx_to_pdf(path, max_retries=3, wait_seconds=1)
                if converted and os.path.exists(converted):
                    pdfs_to_merge.append(converted)
                else:
                    # If PDF conversion failed, try direct text extraction
                    print(f"🔄 PDF conversion failed for {path}, trying direct text extraction...")
                    docx_text = extract_text_from_docx(path)
                    if docx_text:
                        docx_texts.append(docx_text)
                        print(f"✅ Successfully extracted text directly from DOCX: {path}")
                    else:
                        skipped.append((path, "Both PDF conversion and direct text extraction failed"))
            except Exception as e:
                # If PDF conversion failed, try direct text extraction
                print(f"🔄 PDF conversion error for {path}: {e}, trying direct text extraction...")
                docx_text = extract_text_from_docx(path)
                if docx_text:
                    docx_texts.append(docx_text)
                    print(f"✅ Successfully extracted text directly from DOCX: {path}")
                else:
                    skipped.append((path, f"Both PDF conversion and direct text extraction failed: {e}"))
        elif ext == ".pdf":
            if os.path.exists(path) and is_valid_pdf(path):
                pdfs_to_merge.append(path)
            else:
                skipped.append((path, "Invalid or missing PDF"))
        else:
            skipped.append((path, f"Unsupported extension: {ext}"))

    if skipped:
        print("⚠️ Skipped files:")
        for path, reason in skipped:
            print(f"  - {path}: {reason}")

    # Combine all extracted text
    all_text = ""
    
    # Add text from successfully converted PDFs
    if pdfs_to_merge:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        merge_pdfs(pdfs_to_merge, output_path)
        pdf_text = extract_text_from_pdf(output_path)
        if pdf_text:
            all_text += pdf_text + "\n\n"
    
    # Add text from directly extracted DOCX files
    for docx_text in docx_texts:
        if docx_text:
            all_text += docx_text + "\n\n"
    
    # If we have no PDFs to merge but have DOCX text, create a dummy PDF path
    if not pdfs_to_merge and docx_texts:
        # Create a placeholder PDF path for consistency
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        # Write a simple text file as backup (though this won't be a real PDF)
        text_backup_path = output_path.replace('.pdf', '_text_backup.txt')
        with open(text_backup_path, 'w', encoding='utf-8') as f:
            f.write(all_text.strip())
        print(f"📄 Created text backup at: {text_backup_path}")

    if not all_text.strip():
        print("❌ No extractable text found from any files. Skipping submission.")
        return output_path, ""  # Gracefully return empty text for safety

    return output_path, all_text.strip()  # Always return two values

def get_submission_status(sub):
    """
    Classify a Canvas submission by its status.

    Possible return values: "On Time", "Late", "Missing", "Resubmitted"
    """
    if sub.get("attempt") is None or sub.get("missing"):
        return "Missing"
    elif sub.get("late"):
        return "Late"
    elif sub.get("attempt", 1) > 1:
        return "Resubmitted"
    else:
        return "On Time"

